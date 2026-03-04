import nltk
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab', quiet=True)

from PyPDF2 import PdfReader
from docx import Document
from nltk.tokenize import sent_tokenize
from io import BytesIO

def stream_txt(path):
    print(f"[stream_txt] Reading text file: {path}")
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        block_count = 0
        # Read in blocks so it's faster than line by line, but keeps memory low
        while True:
            block = f.read(1024 * 1024) # 1MB blocks
            if not block:
                break
            block_count += 1
            print(f"[stream_txt] Yielding block {block_count} ({len(block)} chars)")
            yield block
    print(f"[stream_txt] Finished reading. Total blocks: {block_count}")

def stream_pdf(path):
    print(f"[stream_pdf] Reading PDF file: {path}")
    try:
        reader = PdfReader(path)
        print(f"[stream_pdf] PDF loaded. Total pages: {len(reader.pages)}")
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                print(f"[stream_pdf] Yielding text from page {i+1} ({len(text)} chars)")
                yield text
            else:
                print(f"[stream_pdf] Page {i+1} has no text, skipping")
        print("[stream_pdf] Finished reading all pages")
    except Exception as e:
        print(f"[stream_pdf] ERROR reading PDF: {e}")
        raise

def stream_docx(path):
    print(f"[stream_docx] Reading DOCX file: {path}")
    try:
        doc = Document(path)
        print(f"[stream_docx] DOCX loaded. Total paragraphs: {len(doc.paragraphs)}")
        para_count = 0
        for p in doc.paragraphs:
            if p.text.strip():
                para_count += 1
                yield p.text
        print(f"[stream_docx] Finished. Yielded {para_count} non-empty paragraphs")
    except Exception as e:
        print(f"[stream_docx] ERROR reading DOCX: {e}")
        raise

def chunk_text_stream(text_stream, sentences_per_chunk=5, overlap=2):
    """
    Given an iterator of text blocks, yield chunks of text incrementally.
    """
    buffer = ""
    sentences = []
    
    for text_part in text_stream:
        buffer += " " + text_part
        new_sentences = sent_tokenize(buffer)
        
        # Keep the last sentence in the buffer because it might be incomplete
        if len(new_sentences) > 1:
            sentences.extend(new_sentences[:-1])
            buffer = new_sentences[-1]
        
        # Yield chunks if we have enough sentences
        while len(sentences) >= sentences_per_chunk:
            chunk = " ".join(sentences[:sentences_per_chunk])
            yield chunk.strip()
            # Slide window
            sentences = sentences[sentences_per_chunk - overlap:]
            
    # Process remaining buffer
    if buffer.strip():
        sentences.extend(sent_tokenize(buffer))
        
    while len(sentences) > 0:
        chunk = " ".join(sentences[:sentences_per_chunk])
        yield chunk.strip()
        if len(sentences) <= sentences_per_chunk:
            break
        sentences = sentences[sentences_per_chunk - overlap:]

def yield_file_chunks(path, sentences_per_chunk=5, overlap=2):
    print(f"[yield_file_chunks] Function called for: {path}")
    if path.lower().endswith(".txt"):
        print("[yield_file_chunks] File type: TXT")
        stream = stream_txt(path)
    elif path.lower().endswith(".pdf"):
        print("[yield_file_chunks] File type: PDF")
        stream = stream_pdf(path)
    elif path.lower().endswith(".docx"):
        print("[yield_file_chunks] File type: DOCX")
        stream = stream_docx(path)
    else:
        print(f"[yield_file_chunks] ERROR: Unsupported file type for {path}")
        return

    print("[yield_file_chunks] Starting chunk_text_stream...")
    chunk_count = 0
    for chunk in chunk_text_stream(stream, sentences_per_chunk, overlap):
        chunk_count += 1
        if chunk_count % 50 == 0:
            print(f"[yield_file_chunks] Yielded {chunk_count} chunks so far...")
        yield chunk
    print(f"[yield_file_chunks] Finished. Total chunks yielded: {chunk_count}")

def stream_uploaded_txt(file_storage):
    raw = file_storage.read()
    text = raw.decode("utf-8", errors="ignore")
    yield text

def stream_uploaded_pdf(file_storage):
    reader = PdfReader(BytesIO(file_storage.read()))
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            yield text

def stream_uploaded_docx(file_storage):
    doc = Document(BytesIO(file_storage.read()))
    for p in doc.paragraphs:
        if p.text.strip():
            yield p.text

def yield_uploaded_file_chunks(file_storage, filename, sentences_per_chunk=5, overlap=2):
    lower_name = filename.lower()
    file_storage.stream.seek(0)

    if lower_name.endswith(".txt"):
        stream = stream_uploaded_txt(file_storage)
    elif lower_name.endswith(".pdf"):
        stream = stream_uploaded_pdf(file_storage)
    elif lower_name.endswith(".docx"):
        stream = stream_uploaded_docx(file_storage)
    else:
        return

    yield from chunk_text_stream(stream, sentences_per_chunk, overlap)
